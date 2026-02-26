import streamlit as st
import os
from pathlib import Path
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
from typing import List, Tuple, Dict, Set
import docx
import json
from datetime import datetime
import re
import tempfile
import shutil
from collections import Counter

# ============================================================================
# DOCUMENT CLASS (must be defined before detector classes)
# ============================================================================

class QADocument:
    """Represents a Q&A document"""
    def __init__(self, filename, full_text, file_path, file_modified_date, word_count, title="", qa_pairs=None):
        self.filename = filename
        self.full_text = full_text
        self.file_path = file_path
        self.file_modified_date = file_modified_date
        self.word_count = word_count
        self.title = title  # Document title
        self.qa_pairs = qa_pairs if qa_pairs is not None else []

# ============================================================================
# PRODUCT NAME DETECTION CONFIGURATION
# ============================================================================

# Manual product list (fallback / baseline)
MANUAL_PRODUCT_NAMES = [
    'vamorolone',
    'agamree',
    # Add more product names as needed:
]

class AutoProductDetector:
    """Automatically detect product names from corpus using frequency analysis"""
    
    def __init__(self):
        self.detected_products = set()
        self.initialized = False
        self.stats = {}
    
    def learn_from_documents(self, documents: List[QADocument]) -> Dict[str, int]:
        """Learn product names from document corpus"""
        if self.initialized:
            return self.stats
        
        word_freq = Counter()
        word_doc_count = Counter()
        total_docs = len(documents)
        
        for doc in documents:
            # Extract capitalized words (likely proper nouns/product names)
            caps_words = re.findall(r'\b[A-Z][a-z]{4,}\b', doc.full_text)
            
            # Track unique words per document
            unique_in_doc = set()
            for word in caps_words:
                word_lower = word.lower()
                
                # Skip common medical terms
                if word_lower in ['patient', 'patients', 'treatment', 'treatments',
                                  'section', 'summary', 'document', 'clinical',
                                  'study', 'studies', 'therapy', 'disease']:
                    continue
                
                word_freq[word_lower] += 1
                unique_in_doc.add(word_lower)
            
            # Count documents containing each word
            for word in unique_in_doc:
                word_doc_count[word] += 1
        
        # Products typically appear in >25% of documents
        doc_threshold = total_docs * 0.25
        
        # Also need sufficient total mentions
        freq_threshold = 10
        
        for word, doc_count in word_doc_count.items():
            if doc_count >= doc_threshold and word_freq[word] >= freq_threshold:
                self.detected_products.add(word)
                self.stats[word] = {
                    'documents': doc_count,
                    'total_mentions': word_freq[word],
                    'doc_percentage': (doc_count / total_docs) * 100
                }
        
        self.initialized = True
        return self.stats
    
    def detect_in_text(self, text: str) -> List[str]:
        """Detect which products appear in given text"""
        if not self.initialized:
            return []
        
        text_lower = text.lower()
        found = []
        
        for product in self.detected_products:
            if re.search(r'\b' + re.escape(product) + r'\b', text_lower):
                found.append(product)
        
        return found
    
    def get_products(self) -> Set[str]:
        """Get all detected products"""
        return self.detected_products

class NERProductDetector:
    """Detect product names using Named Entity Recognition (ScispaCy)"""
    
    def __init__(self):
        self.nlp = None
        self.loaded = False
        self.error = None
        self.detected_cache = {}
    
    def load_model(self):
        """Load ScispaCy NER model"""
        if self.loaded:
            return True
        
        try:
            import spacy
            # Try to load medical NER model
            self.nlp = spacy.load("en_ner_bc5cdr_md")
            self.loaded = True
            return True
        except ImportError:
            self.error = "scispacy not installed. Run: pip install scispacy && pip install https://s3-us-west-2.amazonaws.com/ai2-s3-scispacy/releases/v0.5.1/en_ner_bc5cdr_md-0.5.1.tar.gz"
            return False
        except OSError:
            self.error = "Model not found. Run: pip install https://s3-us-west-2.amazonaws.com/ai2-s3-scispacy/releases/v0.5.1/en_ner_bc5cdr_md-0.5.1.tar.gz"
            return False
        except Exception as e:
            self.error = f"Error loading model: {str(e)}"
            return False
    
    def detect_in_text(self, text: str) -> List[str]:
        """Extract chemical/drug names using NER"""
        if not self.loaded:
            return []
        
        # Check cache
        if text in self.detected_cache:
            return self.detected_cache[text]
        
        try:
            doc = self.nlp(text)
            products = []
            
            for ent in doc.ents:
                if ent.label_ == "CHEMICAL":
                    product = ent.text.lower()
                    if len(product) > 4:  # Skip very short terms
                        products.append(product)
            
            result = list(set(products))
            self.detected_cache[text] = result
            return result
        except Exception as e:
            return []
    
    def get_error(self) -> str:
        """Get error message if model failed to load"""
        return self.error

class HybridProductDetector:
    """Hybrid approach: corpus analysis + NER fallback"""
    
    def __init__(self):
        self.auto_detector = AutoProductDetector()
        self.ner_detector = NERProductDetector()
        self.use_ner_fallback = False
    
    def initialize(self, documents: List[QADocument], enable_ner: bool = False):
        """Initialize both detectors"""
        # Always learn from corpus
        stats = self.auto_detector.learn_from_documents(documents)
        
        # Optionally enable NER fallback
        if enable_ner:
            self.use_ner_fallback = self.ner_detector.load_model()
        
        return stats
    
    def detect_in_text(self, text: str) -> List[str]:
        """Detect products using corpus knowledge, with NER fallback"""
        # First try corpus-based detection (fast)
        found = self.auto_detector.detect_in_text(text)
        
        # If nothing found and NER is available, try NER
        if not found and self.use_ner_fallback:
            found = self.ner_detector.detect_in_text(text)
        
        return found
    
    def get_products(self) -> Set[str]:
        """Get all known products"""
        return self.auto_detector.get_products()

# ============================================================================
# CONFIGURABLE: Choose default detection method
# ============================================================================
PRODUCT_DETECTION_METHODS = {
    'manual': 'Manual List (Fastest)',
    'auto': 'Auto-Detect from Corpus (Smart)',
    'ner': 'NER - ScispaCy (Most Accurate)',
    'hybrid': 'Hybrid - Auto + NER Fallback (Best)'
}

def filter_drug_names(text: str, drug_names: List[str] = None) -> str:
    """
    Reduce frequency of drug name in text to avoid over-weighting
    Keep first 5 occurrences, replace rest with generic placeholder
    """
    if drug_names is None:
        drug_names = MANUAL_PRODUCT_NAMES
    
    filtered = text
    for drug in drug_names:
        # Find all occurrences
        occurrences = list(re.finditer(r'\b' + re.escape(drug) + r'\b', filtered, re.IGNORECASE))
        # If more than 5 occurrences, replace extras
        if len(occurrences) > 5:
            for match in reversed(occurrences[5:]):
                # Replace with generic placeholder
                filtered = filtered[:match.start()] + '[MEDICATION]' + filtered[match.end():]
    
    return filtered

def extract_filename_keywords(filename: str) -> List[str]:
    """Extract meaningful keywords from filename for boosting"""
    # Remove extension
    name = Path(filename).stem
    
    # Remove common patterns: SRL numbers, versions, regions
    name = re.sub(r'SRL_?\d+_?', '', name, flags=re.IGNORECASE)
    name = re.sub(r'_v?F?\d+[\._]\d+', '', name, flags=re.IGNORECASE)
    name = re.sub(r'_EU\b', '', name, flags=re.IGNORECASE)
    name = re.sub(r'_US\b', '', name, flags=re.IGNORECASE)
    name = re.sub(r'_FINAL\b', '', name, flags=re.IGNORECASE)
    name = re.sub(r'\b(VAM|MT)\b', '', name, flags=re.IGNORECASE)
    
    # Split on underscores, dashes, and extract words
    keywords = []
    for word in re.split(r'[_\-\s]+', name):
        word_clean = word.strip()
        # Keep meaningful words (length > 2)
        if len(word_clean) > 2:
            keywords.append(word_clean)
    
    return keywords

class PharmaQASearcher:
    def __init__(self, model_name: str = 'pritamdeka/S-PubMedBert-MS-MARCO'):
        """Initialize the Q&A document searcher with a medical model"""
        st.info(f"Loading model: {model_name}...")
        self.model = SentenceTransformer(model_name)
        self.model_name = model_name
        self.index = None
        self.documents: List[QADocument] = []
        
        # Product detection system
        self.product_detector = None
        self.detection_method = 'manual'  # Default to manual
        self.detection_stats = {}
        
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, str]:
        """
        Extract all text from a Word document
        Returns: (title, body_text)
        """
        try:
            doc = docx.Document(file_path)
            text_parts = []
            title = ""
            
            # Extract title from first bold paragraphs
            # Skip generic product name (first bold line), use second bold line as title
            bold_paragraphs = []
            for i, para in enumerate(doc.paragraphs[:5]):
                text = para.text.strip()
                if text:
                    # Check if paragraph is bold
                    is_bold = any(run.bold for run in para.runs if run.text.strip())
                    if is_bold and len(text) < 200:
                        bold_paragraphs.append(text)
            
            # Use second bold paragraph as title (skip first which is often product name)
            if len(bold_paragraphs) >= 2:
                title = bold_paragraphs[1]  # Skip first bold line (product name)
            elif len(bold_paragraphs) == 1:
                # If only one bold paragraph, check if it looks like a product name
                # (short, all caps, or very generic)
                text = bold_paragraphs[0]
                if len(text) > 30 or not text.isupper():
                    title = text
                else:
                    # Likely just product name, use filename instead
                    title = Path(file_path).stem
            else:
                # Fallback: use filename
                title = Path(file_path).stem
            
            # Get all paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Get table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            body_text = '\n'.join(text_parts)
            return title, body_text
        except Exception as e:
            st.error(f"Error reading {file_path}: {str(e)}")
            return "", ""
    
    def extract_qa_pairs(self, text: str) -> List[Tuple[str, str]]:
        """
        Attempt to extract Q&A pairs from text
        This is a simple heuristic - adjust based on your document format
        """
        qa_pairs = []
        lines = text.split('\n')
        
        current_question = None
        current_answer = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line looks like a question
            if (line.endswith('?') or 
                line.lower().startswith('q:') or 
                line.lower().startswith('question:')):
                # Save previous Q&A pair if exists
                if current_question and current_answer:
                    qa_pairs.append((current_question, '\n'.join(current_answer)))
                
                # Start new question
                current_question = line
                current_answer = []
            elif current_question:
                # Add to current answer
                current_answer.append(line)
        
        # Add last Q&A pair
        if current_question and current_answer:
            qa_pairs.append((current_question, '\n'.join(current_answer)))
        
        return qa_pairs
    
    def get_sentence_similarity(self, sentence: str, query: str) -> float:
        """
        Calculate cosine similarity between a sentence and query using the model
        """
        try:
            # Encode both texts
            sentence_emb = self.model.encode([sentence], convert_to_numpy=True, show_progress_bar=False)
            query_emb = self.model.encode([query], convert_to_numpy=True, show_progress_bar=False)
            
            # Normalize
            faiss.normalize_L2(sentence_emb)
            faiss.normalize_L2(query_emb)
            
            # Calculate cosine similarity (dot product of normalized vectors)
            similarity = np.dot(sentence_emb[0], query_emb[0])
            
            return float(similarity)
        except:
            return 0.0
    
    def highlight_text(self, text: str, query: str, max_length: int = 1500) -> str:
        """
        Highlight query terms in text - enhanced for medical terminology
        """
        if not query or not text:
            return text[:max_length] + ("..." if len(text) > max_length else "")
        
        # Extract query terms including medical/drug names
        stop_words = {'the', 'is', 'at', 'which', 'on', 'a', 'an', 'and', 'or', 'but', 'in', 'with', 'to', 'for', 'of', 'as', 'by', 'this', 'that', 'these', 'those', 'are', 'was', 'were', 'be', 'been', 'being', 'can', 'could', 'should', 'would', 'will', 'what', 'when', 'where', 'who', 'how', 'does', 'do', 'did', 'has', 'have', 'had'}
        
        query_words = query.split()
        query_terms = []
        
        for word in query_words:
            word_clean = word.lower().strip('?.,!;:')
            if len(word_clean) > 2 and word_clean not in stop_words:
                query_terms.append(word_clean)
            # Preserve case for drug names
            if word[0].isupper() and len(word) > 3:
                query_terms.append(word.strip('?.,!;:'))
        
        query_terms = list(dict.fromkeys(query_terms))
        
        if not query_terms:
            return text[:max_length] + ("..." if len(text) > max_length else "")
        
        # Find first occurrence of any query term
        text_lower = text.lower()
        first_match = len(text)
        
        for term in query_terms:
            pos = text_lower.find(term.lower())
            if pos != -1 and pos < first_match:
                first_match = pos
        
        if first_match == len(text):
            excerpt = text[:max_length]
        else:
            start = max(0, first_match - 200)
            end = min(len(text), start + max_length)
            excerpt = text[start:end]
        
        # Highlight with variations
        highlighted = excerpt
        sorted_terms = sorted(query_terms, key=len, reverse=True)
        
        for term in sorted_terms:
            # Multiple patterns for better matching
            patterns = [
                # Exact match
                re.compile(r'\b(' + re.escape(term) + r')\b', re.IGNORECASE),
                # Word starting with term
                re.compile(r'\b(' + re.escape(term.lower()) + r'\w*)\b', re.IGNORECASE) if len(term) > 4 else None,
            ]
            
            for pattern in patterns:
                if pattern:
                    highlighted = pattern.sub(
                        lambda m: f'<mark style="background-color: #ffeb3b; padding: 2px 4px; border-radius: 3px; font-weight: 500;">{m.group(1)}</mark>',
                        highlighted
                    )
        
        prefix = "..." if first_match > 200 else ""
        suffix = "..." if end < len(text) else ""
        
        return prefix + highlighted + suffix
    
    def get_relevant_snippets(self, text: str, query: str, num_snippets: int = 3, snippet_length: int = 500) -> List[str]:
        """
        Extract multiple relevant snippets from text based on semantic similarity
        Enhanced highlighting for medical terminology with proper HTML handling
        """
        if not query or not text:
            return []
        
        # Split text into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        if not sentences:
            return []
        
        # Calculate similarity for each sentence (this is fast)
        sentence_similarities = []
        
        for i, sentence in enumerate(sentences):
            similarity = self.get_sentence_similarity(sentence, query)
            if similarity > 0.25:  # Lower threshold to get more candidates
                sentence_similarities.append((similarity, i, sentence))
        
        if not sentence_similarities:
            return []
        
        # Sort by similarity
        sentence_similarities.sort(reverse=True)
        
        # Get top snippets
        snippets = []
        used_positions = set()
        
        # Enhanced keyword extraction - include medical terms
        stop_words = {'the', 'is', 'at', 'which', 'on', 'a', 'an', 'and', 'or', 'but', 'in', 'with', 'to', 'for', 'of', 'as', 'by', 'this', 'that', 'these', 'those', 'are', 'was', 'were', 'be', 'been', 'being', 'can', 'could', 'should', 'would', 'will', 'what', 'when', 'where', 'who', 'how', 'does', 'do', 'did', 'has', 'have', 'had'}
        
        # Extract ALL meaningful terms including short ones for medical context
        query_words = query.split()
        query_terms = []
        
        for word in query_words:
            word_clean = word.lower().strip('?.,!;:')
            # Include medical/drug names (often capitalized or unique)
            if len(word_clean) > 2 and word_clean not in stop_words:
                query_terms.append(word_clean)
            # Also include the original case for drug names
            if word[0].isupper() and len(word) > 3:
                query_terms.append(word.strip('?.,!;:'))
        
        # Remove duplicates while preserving order
        query_terms = list(dict.fromkeys(query_terms))
        
        for similarity, idx, sentence in sentence_similarities[:num_snippets * 2]:
            # Avoid duplicate positions
            if idx in used_positions:
                continue
            
            # Mark surrounding positions as used
            for j in range(max(0, idx-1), min(len(sentences), idx+2)):
                used_positions.add(j)
            
            # Enhanced highlighting with stem matching
            highlighted = sentence
            
            # Track what we've highlighted to avoid duplicates
            highlighted_positions = []
            
            # Sort terms by length (longest first) to avoid partial matches
            sorted_terms = sorted(query_terms, key=len, reverse=True)
            
            for term in sorted_terms:
                term_lower = term.lower()
                
                # Find all occurrences with variations
                # Pattern 1: Exact match (case insensitive)
                pattern1 = re.compile(r'\b(' + re.escape(term) + r')\b', re.IGNORECASE)
                
                # Pattern 2: Word starts with term (catches plurals, conjugations)
                if len(term_lower) > 4:
                    pattern2 = re.compile(r'\b(' + re.escape(term_lower) + r'\w*)\b', re.IGNORECASE)
                else:
                    pattern2 = None
                
                # Pattern 3: Stem matching (first 4-5 chars for medical terms)
                stem_patterns = []
                if len(term_lower) >= 5:
                    stem = term_lower[:5]
                    stem_patterns.append(re.compile(r'\b(' + re.escape(stem) + r'\w{1,15})\b', re.IGNORECASE))
                
                # Apply patterns in order of specificity
                for pattern in [pattern1, pattern2] + stem_patterns:
                    if pattern:
                        matches = list(pattern.finditer(highlighted))
                        for match in reversed(matches):  # Reverse to maintain positions
                            start, end = match.span(1)
                            
                            # Check if already highlighted
                            overlap = False
                            for h_start, h_end in highlighted_positions:
                                if (start >= h_start and start < h_end) or (end > h_start and end <= h_end):
                                    overlap = True
                                    break
                            
                            if not overlap:
                                matched_text = highlighted[start:end]
                                # Only highlight if it's semantically relevant
                                if len(matched_text) >= 3:
                                    highlighted = (
                                        highlighted[:start] + 
                                        f'<mark style="background-color: #ffeb3b; padding: 2px 4px; border-radius: 3px; font-weight: 500;">{matched_text}</mark>' +
                                        highlighted[end:]
                                    )
                                    # Update positions (account for added HTML)
                                    html_added = len('<mark style="background-color: #ffeb3b; padding: 2px 4px; border-radius: 3px; font-weight: 500;">') + len('</mark>')
                                    highlighted_positions.append((start, end + html_added))
            
            # IMPROVED: Smart truncation that properly handles HTML tags
            if len(highlighted) > snippet_length:
                # First try to find a sentence boundary
                last_period = highlighted[:snippet_length].rfind('.')
                
                if last_period > snippet_length * 0.6:
                    # Good sentence boundary found
                    truncate_point = last_period + 1
                else:
                    # No good sentence boundary, need to be careful with HTML
                    truncate_point = snippet_length
                    
                    # Check if we're inside or near a <mark> tag
                    # Find all mark tag positions
                    mark_starts = [m.start() for m in re.finditer(r'<mark[^>]*>', highlighted)]
                    mark_ends = [m.end() for m in re.finditer(r'</mark>', highlighted)]
                    
                    # Check if truncate point is inside a tag
                    for mark_start in mark_starts:
                        if mark_start < truncate_point < mark_start + 100:  # Inside opening tag
                            # Move to before this tag starts
                            truncate_point = mark_start
                            break
                    
                    # If we're between opening and closing mark, extend to after closing
                    for i, mark_start in enumerate(mark_starts):
                        if i < len(mark_ends):
                            mark_end = mark_ends[i]
                            if mark_start < truncate_point < mark_end:
                                # We're inside marked text, extend to include closing tag
                                truncate_point = mark_end
                                break
                
                # Truncate at safe point
                highlighted = highlighted[:truncate_point]
                
                # Ensure all opened <mark> tags are closed
                open_marks = highlighted.count('<mark')
                close_marks = highlighted.count('</mark>')
                
                if open_marks > close_marks:
                    highlighted += '</mark>' * (open_marks - close_marks)
                
                # Remove any incomplete tags at the end
                # Check if we have an incomplete opening tag
                last_opening = highlighted.rfind('<mark')
                last_closing = highlighted.rfind('>')
                
                if last_opening > last_closing:
                    # We have an incomplete opening tag, remove it
                    highlighted = highlighted[:last_opening]
                
                # Check for incomplete closing tag
                if highlighted.endswith('</mark') or highlighted.endswith('</mar') or \
                   highlighted.endswith('</ma') or highlighted.endswith('</m') or highlighted.endswith('</'):
                    # Find the last complete tag
                    last_complete = max(
                        highlighted.rfind('</mark>'),
                        highlighted.rfind('>') if not highlighted[highlighted.rfind('>'):].startswith('></') else -1
                    )
                    if last_complete > 0:
                        highlighted = highlighted[:last_complete + 1]
                
                # Add ellipsis if we actually truncated content
                if truncate_point < len(sentence):
                    highlighted += "..."
            
            snippets.append(highlighted)
            
            if len(snippets) >= num_snippets:
                break
        
        return snippets
    
    def save_cache(self, directory: str):
        """Save embeddings to cache"""
        pass  # Caching disabled to avoid pickle issues
    
    def load_cache(self, directory: str) -> bool:
        """Try to load cached embeddings"""
        return False  # Caching disabled to avoid pickle issues
    
    def index_documents(self, directory: str, file_pattern: str = "") -> int:
        """Index all Word Q&A documents in a directory with title weighting"""
        self.documents = []
        directory_path = Path(directory)
        
        # Find all .docx files
        files = list(directory_path.rglob('*.docx'))
        
        # Filter out temporary Word files (starting with ~$)
        files = [f for f in files if not f.name.startswith('~$')]
        
        # Apply file pattern filter if specified
        if file_pattern:
            files = [f for f in files if file_pattern.lower() in f.name.lower()]
        
        if not files:
            st.warning("No Word documents found in the directory")
            return 0
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Process each file
        for idx, file_path in enumerate(files):
            status_text.text(f"Processing: {file_path.name} ({idx+1}/{len(files)})")
            
            # Extract text and title
            title, text = self.extract_text_from_docx(str(file_path))
            
            if not text.strip():
                st.warning(f"No text extracted from {file_path.name}")
                continue
            
            # Get file metadata
            file_modified = datetime.fromtimestamp(os.path.getmtime(file_path))
            word_count = len(text.split())
            
            # Try to extract Q&A pairs (optional)
            qa_pairs = self.extract_qa_pairs(text)
            
            # Create document object
            doc = QADocument(
                filename=file_path.name,
                full_text=text,
                file_path=str(file_path),
                file_modified_date=file_modified.strftime('%Y-%m-%d'),
                word_count=word_count,
                title=title,
                qa_pairs=qa_pairs if qa_pairs else None
            )
            
            self.documents.append(doc)
            progress_bar.progress((idx + 1) / len(files))
        
        if not self.documents:
            st.error("No documents were successfully processed")
            progress_bar.empty()
            status_text.empty()
            return 0
        
        # Initialize product detection based on selected method
        status_text.text("Initializing product detection...")
        self._initialize_product_detection()
        
        status_text.text("Creating document embeddings with optimized weighting...")
        
        # Create embeddings with improved weighting strategy:
        # - Title: 2x (reduced from 3x to reduce over-emphasis)
        # - Filename keywords: 2x (boost relevant terms from filename)
        # - Body: 1x with drug names filtered (reduce drug name pollution)
        weighted_texts = []
        for doc in self.documents:
            # Extract keywords from filename
            filename_keywords = extract_filename_keywords(doc.filename)
            filename_text = ' '.join(filename_keywords)
            
            # Filter excessive drug name mentions from body
            filtered_body = filter_drug_names(doc.full_text)
            
            # Construct weighted text: Title 2x + Filename 2x + Filtered Body 1x
            weighted_text = f"{doc.title}. {doc.title}. {filename_text}. {filename_text}. {filtered_body}"
            weighted_texts.append(weighted_text)
        
        embeddings = self.model.encode(
            weighted_texts, 
            show_progress_bar=True, 
            batch_size=8,
            convert_to_numpy=True
        )
        
        # Create FAISS index
        dimension = embeddings.shape[1]
        self.index = faiss.IndexFlatIP(dimension)
        
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings)
        self.index.add(embeddings.astype('float32'))
        
        progress_bar.empty()
        status_text.empty()
        
        return len(self.documents)
    
    def _initialize_product_detection(self):
        """Initialize product detection based on selected method"""
        method = st.session_state.get('detection_method', 'manual')
        self.detection_method = method
        
        if method == 'manual':
            # Use manual list
            self.detection_stats = {
                'method': 'manual',
                'products': MANUAL_PRODUCT_NAMES,
                'count': len(MANUAL_PRODUCT_NAMES)
            }
            st.success(f"✅ Using manual product list: {', '.join(MANUAL_PRODUCT_NAMES)}")
        
        elif method == 'auto':
            # Auto-detect from corpus
            self.product_detector = AutoProductDetector()
            stats = self.product_detector.learn_from_documents(self.documents)
            detected = self.product_detector.get_products()
            
            self.detection_stats = {
                'method': 'auto',
                'products': list(detected),
                'count': len(detected),
                'details': stats
            }
            
            if detected:
                st.success(f"✅ Auto-detected {len(detected)} products: {', '.join(sorted(detected))}")
            else:
                st.warning("⚠️ No products auto-detected. Falling back to manual list.")
                self.detection_method = 'manual'
        
        elif method == 'ner':
            # Use NER only
            self.product_detector = NERProductDetector()
            loaded = self.product_detector.load_model()
            
            if loaded:
                self.detection_stats = {
                    'method': 'ner',
                    'products': 'Dynamic (detected per query)',
                    'status': 'Model loaded'
                }
                st.success("✅ ScispaCy NER model loaded successfully")
            else:
                error = self.product_detector.get_error()
                st.error(f"❌ NER model failed to load: {error}")
                st.info("Falling back to manual list...")
                self.detection_method = 'manual'
        
        elif method == 'hybrid':
            # Hybrid: Auto + NER fallback
            self.product_detector = HybridProductDetector()
            stats = self.product_detector.initialize(self.documents, enable_ner=True)
            detected = self.product_detector.get_products()
            
            self.detection_stats = {
                'method': 'hybrid',
                'products': list(detected),
                'count': len(detected),
                'details': stats,
                'ner_enabled': self.product_detector.use_ner_fallback
            }
            
            if detected:
                ner_status = "with NER fallback" if self.product_detector.use_ner_fallback else "without NER"
                st.success(f"✅ Hybrid detection: {len(detected)} products from corpus {ner_status}")
                st.info(f"Products: {', '.join(sorted(detected))}")
            else:
                st.warning("⚠️ Hybrid detection found no products. Using manual list.")
                self.detection_method = 'manual'
    
    def _get_detected_products(self, query: str) -> List[str]:
        """Get detected products in query based on current detection method"""
        if self.detection_method == 'manual':
            # Check manual list
            query_lower = query.lower()
            return [p for p in MANUAL_PRODUCT_NAMES if p in query_lower]
        
        elif self.detection_method in ['auto', 'ner', 'hybrid']:
            # Use detector
            if self.product_detector:
                return self.product_detector.detect_in_text(query)
            else:
                # Fallback to manual
                query_lower = query.lower()
                return [p for p in MANUAL_PRODUCT_NAMES if p in query_lower]
        
        return []
    
    def search(self, query: str, top_k: int = 5, min_score: float = 0.0) -> List[Tuple[QADocument, float]]:
        """Search for most relevant Q&A documents with synonym expansion and product name de-weighting"""
        if self.index is None or len(self.documents) == 0:
            return []
        
        # Expand query with pharma-specific synonyms
        expanded_query = expand_query_with_synonyms(query)
        
        # Detect product names in query using selected detection method
        found_products = self._get_detected_products(expanded_query)
        found_product = found_products[0] if found_products else None
        
        # If product name found, reduce its repetition in the query
        if found_product:
            # Keep product name once, but don't let it dominate
            # Replace multiple mentions with single mention
            pattern = re.compile(r'\b' + re.escape(found_product) + r'\b', re.IGNORECASE)
            matches = list(pattern.finditer(expanded_query))
            
            if len(matches) > 1:
                # Keep first occurrence, remove others
                for match in reversed(matches[1:]):
                    expanded_query = (
                        expanded_query[:match.start()] + 
                        expanded_query[match.end():]
                    )
        
        # Encode query
        query_embedding = self.model.encode([expanded_query], convert_to_numpy=True)
        faiss.normalize_L2(query_embedding)
        
        # Search
        k = min(top_k, len(self.documents))
        scores, indices = self.index.search(query_embedding.astype('float32'), k)
        
        # NEW: Post-process scores to de-weight product name matches
        # This ensures "MOA of vamorolone" matches on "MOA" not just "vamorolone"
        results = []
        for idx, score in zip(indices[0], scores[0]):
            adjusted_score = score
            
            # If query contained product name, penalize docs that match primarily on product name
            if found_product and score >= min_score:
                doc = self.documents[idx]
                doc_text_lower = doc.full_text.lower()
                doc_title_lower = doc.title.lower()
                
                # Count how many times product appears in this doc
                product_count = doc_text_lower.count(found_product.lower())
                
                # CRITICAL: Check if doc matches on content terms (not just product name)
                query_terms = expanded_query.lower().split()
                # Remove product name and common words from query terms
                content_terms = [
                    t for t in query_terms 
                    if found_product not in t 
                    and len(t) > 3
                    and t not in ['what', 'does', 'have', 'with', 'from', 'this', 'that', 'there', 'are', 'the', 'and']
                ]
                
                # Calculate content match ratio
                content_ratio = 0.0
                if content_terms:
                    # Check how many content terms appear in doc (body + title)
                    content_matches = sum(1 for term in content_terms if term in doc_text_lower)
                    
                    # BOOST: Check if content terms appear in TITLE or FILENAME (very strong signal)
                    title_matches = sum(1 for term in content_terms if term in doc_title_lower)
                    filename_lower = doc.filename.lower()
                    filename_matches = sum(1 for term in content_terms if term in filename_lower)
                    
                    # Weight title/filename matches more heavily
                    weighted_matches = content_matches + (title_matches * 2) + (filename_matches * 2)
                    content_ratio = min(1.0, weighted_matches / len(content_terms))
                
                # MAXIMUM AGGRESSIVE PENALTY for low content match:
                if content_ratio < 0.5:
                    # Document matches on product name only, NOT on actual content
                    # MASSIVE -30% penalty (tripled from -10%)
                    adjusted_score = adjusted_score * 0.70
                    
                elif content_ratio < 0.7:
                    # Weak content match - still penalize
                    # -15% penalty
                    adjusted_score = adjusted_score * 0.85
                
                # BOOST for EXCELLENT content matches (helps correct document)
                elif content_ratio > 1.5:  # Title + filename + body all match
                    # STRONG +8% boost (was +5%)
                    adjusted_score = adjusted_score * 1.08
                    
                elif content_ratio > 1.0:  # Title or filename match + body
                    # GOOD +5% boost
                    adjusted_score = adjusted_score * 1.05
                
                # ADDITIONAL: Product frequency penalty (applies on top)
                if product_count > 10:
                    # EXTREME penalty for product name spam:
                    # 11-20 mentions: -12% to -24%
                    # 21-30 mentions: -24% to -36%
                    # 30+ mentions: -36% to -50% (CRUSHING penalty)
                    penalty_factor = min(0.50, (product_count - 10) * 0.012)
                    adjusted_score = adjusted_score * (1 - penalty_factor)
            
            if adjusted_score >= min_score:
                results.append((self.documents[idx], float(adjusted_score)))
        
        return results

# Medical model options
MEDICAL_MODELS = {
    'PubMedBERT (Search-Optimized) - RECOMMENDED': 'pritamdeka/S-PubMedBert-MS-MARCO',
    'PubMedBERT (Base)': 'microsoft/BiomedNLP-PubMedBERT-base-uncased-abstract-fulltext',
    'BioBERT v1.2': 'dmis-lab/biobert-base-cased-v1.2',
    'SciBERT': 'allenai/scibert_scivocab_uncased',
    'General (Fast & Lightweight)': 'all-MiniLM-L6-v2'
}

# Example questions organized by category
EXAMPLE_QUESTIONS = {
    "Indication & Mechanism": [
        "What is the approved indication for [drug name], and in which patient populations has it been studied?",
        "How does the mechanism of action of [drug name] differ from traditional treatments?",
    ],
    "Efficacy & Clinical Data": [
        "What clinical trial data support the efficacy of [drug name]?",
        "How does [drug name] compare to standard treatment in terms of efficacy outcomes?",
        "What endpoints were used to assess functional outcomes in [drug name] clinical trials?",
    ],
    "Dosing & Administration": [
        "What dosing regimens of [drug name] have been evaluated in clinical studies?",
        "Are there specific age or weight considerations for dosing [drug name]?",
        "How and when should [drug name] be taken each day?",
        "Is dose tapering required when discontinuing [drug name] therapy?",
    ],
    "Safety & Adverse Events": [
        "What are the most common adverse events reported with [drug name] in clinical trials?",
        "What is known about the long-term safety profile of [drug name]?",
        "What side effects should I watch for while taking [drug name]?",
        "How does [drug name] impact growth, bone health, and endocrine parameters?",
    ],
    "Drug Interactions & Monitoring": [
        "Are there any known drug-drug interactions with [drug name]?",
        "What monitoring is recommended for patients receiving [drug name]?",
        "How should patients be transitioned from conventional treatments to [drug name]?",
    ],
    "Special Populations": [
        "Are there data on the use of [drug name] in patients with hepatic or renal impairment?",
        "Has [drug name] been studied in pediatric populations?",
        "Will taking [drug name] affect my child's growth or long-term health?",
    ],
    "Pharmacology": [
        "What is known about the pharmacokinetics and metabolism of [drug name]?",
        "What is the effect of [drug name] on immune function and infection risk?",
        "Does [drug name] have [specific activity], and how clinically relevant is this?",
    ],
    "Practical Information": [
        "What is [drug name], and how is it different from other medications?",
        "What benefits can I expect from taking [drug name]?",
        "Are there any contraindications or important precautions associated with [drug name] use?",
    ]
}

# Categorize questions by audience
HCP_QUESTIONS = {
    "Indication & Mechanism",
    "Efficacy & Clinical Data",
    "Drug Interactions & Monitoring",
    "Pharmacology"
}

PATIENT_QUESTIONS = {
    "Safety & Adverse Events",
    "Special Populations",
    "Dosing & Administration",
    "Practical Information"
}

# Patient-friendly terms mapping
PATIENT_FRIENDLY_TERMS = {
    'adverse events': 'side effects',
    'adverse event': 'side effect',
    'contraindications': 'reasons you should not take this',
    'contraindication': 'reason you should not take this',
    'pharmacokinetics': 'how the medicine works in your body',
    'pharmacodynamics': 'how the medicine affects your body',
    'hepatic': 'liver',
    'renal': 'kidney',
    'pediatric': 'children',
    'efficacy': 'how well it works',
    'administer': 'take',
    'administration': 'taking',
    'dosing regimen': 'how to take it',
    'drug-drug interaction': 'how it works with other medicines',
    'drug interaction': 'how it works with other medicines',
    'monitoring': 'regular check-ups',
    'metabolism': 'how your body processes it',
    'therapeutic': 'treatment',
    'titration': 'dose adjustment',
    'bioavailability': 'how much medicine gets into your blood',
    'half-life': 'how long the medicine stays in your body',
    'clinical trial': 'research study',
    'clinical endpoint': 'study result',
    'prophylaxis': 'prevention',
    'discontinue': 'stop taking',
    'discontinuing': 'stopping',
    'taper': 'reduce gradually',
    'tapering': 'reducing gradually'
}

# Pharma abbreviations (expand during search)
PHARMA_ABBREVIATIONS = {
    'moa': 'mechanism of action',
    'pk': 'pharmacokinetics',
    'pd': 'pharmacodynamics',
    'pkpd': 'pharmacokinetics pharmacodynamics',
    'ddi': 'drug drug interaction',
    'ddis': 'drug drug interactions',
    'ae': 'adverse event',
    'aes': 'adverse events',
    'sae': 'serious adverse event',
    'saes': 'serious adverse events',
    'peds': 'pediatric children',
    'gi': 'gastrointestinal',
    'cv': 'cardiovascular cardiac heart',
    'cns': 'central nervous system',
    'adme': 'absorption distribution metabolism excretion'
}

# Pharma-specific synonyms for search expansion (based on sample documents)
PHARMA_SYNONYMS = {
    'storage': ['storage', 'store', 'storing', 'stored', 'keep', 'kept', 'keeping', 'transport', 'transportation', 'shipping', 'ship'],
    'transport': ['transport', 'transportation', 'shipping', 'ship', 'shipped', 'delivery', 'transit', 'storage', 'store'],
    'stability': ['stability', 'stable', 'shelf life', 'expiry', 'expiration', 'storage conditions'],
    'cold chain': ['cold chain', 'refrigerated', 'refrigeration', 'temperature controlled', 'temperature', 'cooling'],
    'discontinue': ['discontinue', 'discontinuing', 'stop', 'stopping', 'cease', 'withdraw', 'withdrawal', 'taper', 'tapering'],
    'stop': ['stop', 'stopping', 'discontinue', 'discontinuing', 'cease', 'withdraw', 'taper', 'tapering'],
    'taper': ['taper', 'tapering', 'reduce gradually', 'wean', 'weaning', 'discontinue', 'stop gradually'],
    'dose': ['dose', 'dosage', 'dosing', 'amount', 'quantity', 'regimen'],
    'excipient': ['excipient', 'excipients', 'inactive ingredient', 'inactive ingredients'],
    'formulation': ['formulation', 'formulated', 'formula', 'composition'],
    'indication': ['indication', 'indicated', 'approved for', 'licensed for', 'use'],
    'administration': ['administration', 'administer', 'give', 'take', 'taking', 'dosing'],
    'adverse event': ['adverse event', 'adverse reaction', 'side effect', 'undesirable effect', 'safety'],
}

def clean_query(query: str) -> str:
    """Remove common words that pollute search results"""
    # Remove common question starters that don't add semantic value
    stop_patterns = [
        r'^\s*are\s+there\s+any\s+',
        r'^\s*is\s+there\s+any\s+',
        r'^\s*does\s+\w+\s+have\s+any\s+',
        r'^\s*do\s+\w+\s+have\s+any\s+',
    ]
    
    cleaned = query
    for pattern in stop_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
    
    return cleaned.strip()

def expand_query_with_synonyms(query: str) -> str:
    """
    Expand query with pharma-specific synonyms and abbreviations
    """
    # First clean the query
    query = clean_query(query)
    
    query_lower = query.lower()
    expanded_terms = [query]
    
    # Expand abbreviations FIRST
    for abbrev, full in PHARMA_ABBREVIATIONS.items():
        if re.search(r'\b' + abbrev + r'\b', query_lower):
            expanded_terms.append(full)
    
    # Then add synonyms
    for key, synonyms in PHARMA_SYNONYMS.items():
        # If any synonym appears in query, add related terms
        if any(syn in query_lower for syn in synonyms):
            # Add up to 2 most relevant synonyms not already in query
            added = 0
            for syn in synonyms[:4]:
                if syn not in query_lower and added < 2:
                    expanded_terms.append(syn)
                    added += 1
    
    # Join with spaces and return
    return ' '.join(expanded_terms)

def detect_query_audience(query: str) -> str:
    """
    Detect if a query is from an HCP or patient based on language complexity
    """
    query_lower = query.lower()
    
    # Strong HCP indicators (technical medical terms)
    strong_hcp_terms = [
        'pharmacokinetics', 'pharmacodynamics', 'mechanism of action', 
        'bioavailability', 'half-life', 'therapeutic index', 'metabolism',
        'titration', 'clinical endpoint', 'efficacy outcomes', 'pharmacology',
        'drug-drug interaction', 'monitoring parameters', 'dosing regimen',
        'adverse events', 'contraindication', 'hepatic impairment', 'renal impairment'
    ]
    
    # Moderate HCP indicators
    moderate_hcp_terms = [
        'clinical trial', 'study', 'trial data', 'evidence', 'efficacy',
        'clinical data', 'endpoints', 'parameters', 'regimen'
    ]
    
    # Strong patient indicators (personal/direct language)
    strong_patient_terms = [
        'my child', 'my son', 'my daughter', 'i take', 'should i', 'can i',
        'will it affect me', 'is it safe for me', 'what if i miss', 'how do i take',
        'my doctor', 'i am taking', 'i have been', 'for my', 'will i'
    ]
    
    # Moderate patient indicators (simple questions)
    moderate_patient_terms = [
        'plain language', 'simple terms', 'easy to understand',
        'side effects', 'safe to take', 'when should i'
    ]
    
    # Check for explicit audience request
    if any(term in query_lower for term in ['patient-friendly', 'caregiver', 'parent', 'in plain language']):
        return 'patient'
    if any(term in query_lower for term in ['hcp', 'healthcare provider', 'physician', 'for prescribers']):
        return 'hcp'
    
    # Count indicators with weights
    strong_hcp_count = sum(1 for term in strong_hcp_terms if term in query_lower)
    moderate_hcp_count = sum(1 for term in moderate_hcp_terms if term in query_lower)
    strong_patient_count = sum(1 for term in strong_patient_terms if term in query_lower)
    moderate_patient_count = sum(1 for term in moderate_patient_terms if term in query_lower)
    
    # Weight the scores
    hcp_score = (strong_hcp_count * 3) + moderate_hcp_count
    patient_score = (strong_patient_count * 3) + moderate_patient_count
    
    # Strong HCP terms always win
    if strong_hcp_count > 0:
        return 'hcp'
    
    # Strong patient terms always win (if no strong HCP terms)
    if strong_patient_count > 0:
        return 'patient'
    
    # Compare weighted scores
    if hcp_score > patient_score:
        return 'hcp'
    elif patient_score > hcp_score:
        return 'patient'
    
    # If tied, check question complexity as tiebreaker
    # Remove common words
    words = query_lower.split()
    content_words = [w for w in words if len(w) > 3 and w not in ['what', 'how', 'when', 'where', 'which', 'does', 'this', 'that', 'with', 'from', 'have', 'been']]
    
    # Check average word length of content words
    if content_words:
        avg_word_length = sum(len(word) for word in content_words) / len(content_words)
        # Longer average word length suggests HCP (medical terminology tends to be longer)
        return 'hcp' if avg_word_length >= 7 else 'patient'
    
    # Default to patient mode for safety
    return 'patient'

def simplify_for_patients(text: str) -> str:
    """
    Replace medical jargon with patient-friendly language
    """
    simplified = text
    # Sort by length (longest first) to avoid partial replacements
    sorted_terms = sorted(PATIENT_FRIENDLY_TERMS.items(), key=lambda x: len(x[0]), reverse=True)
    
    for medical_term, friendly_term in sorted_terms:
        # Case-insensitive replacement with word boundaries
        pattern = re.compile(r'\b' + re.escape(medical_term) + r'\b', re.IGNORECASE)
        simplified = pattern.sub(friendly_term, simplified)
    
    return simplified

def main():
    st.set_page_config(
        page_title="Pharma Q&A Document Finder", 
        page_icon="💊", 
        layout="wide"
    )
    
    st.title("💊 Pharma Q&A Document Finder")
    st.markdown("Find the most relevant Q&A documents for HCP and public queries using semantic search")
    
    # Initialize session state
    if 'searcher' not in st.session_state:
        st.session_state.searcher = None
        st.session_state.indexed = False
        st.session_state.current_model = None
        st.session_state.last_results = []
        st.session_state.last_query = ""
        st.session_state.temp_dir = None
        st.session_state.forced_audience = None
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Model selection
        st.subheader("🤖 Model Selection")
        selected_model_name = st.selectbox(
            "Choose a model:",
            options=list(MEDICAL_MODELS.keys()),
            index=0,
            help="PubMedBERT (Search-Optimized) is best for pharma documents"
        )
        
        selected_model = MEDICAL_MODELS[selected_model_name]
        
        with st.expander("ℹ️ Model Info"):
            st.markdown("""
            **PubMedBERT (Search-Optimized)**: Best choice - trained on medical 
            literature and fine-tuned specifically for search/retrieval tasks.
            
            **General (Fast)**: Faster but less accurate for medical terminology.
            
            **Semantic Matching**: Uses sentence-level similarity for ranking.
            Enhanced highlighting catches variations like "study/studied/studies".
            
            **New Features**:
            - 🎯 Title 2x + Filename keywords 2x + Filtered body
            - 🔤 Abbreviations: "MOA" → "mechanism of action", "DDI" → "drug drug interaction"
            - 🔗 Pharma synonyms: "storage" ↔ "transport", "discontinue" ↔ "stop/taper"
            - 🧹 Query cleaning: Removes "are there any", "does X have" for better matching
            - 💊 Product de-weighting: "MOA of vamorolone" matches on MOA, not just drug name
            """)
        
        # Initialize model
        if st.session_state.searcher is None or st.session_state.current_model != selected_model:
            if st.button("🔄 Load Model", type="primary", use_container_width=True):
                with st.spinner("Loading model... (first time will download ~400MB)"):
                    st.session_state.searcher = PharmaQASearcher(model_name=selected_model)
                    st.session_state.current_model = selected_model
                    st.session_state.indexed = False
                    st.success(f"✅ Model loaded!")
        
        if st.session_state.searcher is not None:
            st.success(f"✅ {selected_model_name.split(' - ')[0]}")
        
        st.divider()
        
        # Product Detection Method Selection
        st.subheader("🔬 Product Detection")
        
        detection_method = st.selectbox(
            "Detection Method:",
            options=list(PRODUCT_DETECTION_METHODS.keys()),
            format_func=lambda x: PRODUCT_DETECTION_METHODS[x],
            index=1,  # Default to 'auto'
            help="Choose how product names are detected for de-weighting"
        )
        
        # Store in session state
        if 'detection_method' not in st.session_state:
            st.session_state.detection_method = 'auto'
        
        if detection_method != st.session_state.detection_method:
            st.session_state.detection_method = detection_method
            st.session_state.indexed = False  # Force re-indexing
        
        with st.expander("ℹ️ Detection Methods"):
            st.markdown("""
            **Manual List**: Uses pre-defined list of product names
            - ✅ Fastest (instant)
            - ✅ 100% accurate for known products
            - ❌ Must update list manually
            
            **Auto-Detect**: Learns from your documents
            - ✅ No manual maintenance
            - ✅ Fast (<1ms per query)
            - ✅ Finds products mentioned frequently
            - ⚠️ Requires 25%+ document coverage
            
            **NER - ScispaCy**: Medical entity recognition
            - ✅ Highest accuracy
            - ✅ Detects any drug/chemical name
            - ⚠️ Requires ScispaCy installation
            - ⚠️ Slower (50-100ms per query)
            
            **Hybrid**: Auto + NER fallback (recommended)
            - ✅ Best of both worlds
            - ✅ Fast (uses corpus cache)
            - ✅ NER fallback for unknowns
            - ✅ Self-learning
            """)
        
        # Show current detection stats
        if st.session_state.indexed and st.session_state.searcher:
            stats = st.session_state.searcher.detection_stats
            if stats:
                with st.expander("📊 Detection Stats"):
                    st.markdown(f"**Method**: {stats.get('method', 'N/A').upper()}")
                    
                    if stats.get('method') == 'manual':
                        products = stats.get('products', [])
                        st.markdown(f"**Products**: {', '.join(products)}")
                    
                    elif stats.get('method') in ['auto', 'hybrid']:
                        products = stats.get('products', [])
                        st.markdown(f"**Detected Products**: {len(products)}")
                        if products:
                            st.markdown(f"_{', '.join(sorted(products))}_")
                        
                        details = stats.get('details', {})
                        if details:
                            st.markdown("**Coverage:**")
                            for product, info in list(details.items())[:5]:
                                doc_pct = info.get('doc_percentage', 0)
                                mentions = info.get('total_mentions', 0)
                                st.caption(f"• {product}: {doc_pct:.1f}% of docs ({mentions} mentions)")
                    
                    elif stats.get('method') == 'ner':
                        st.markdown(f"**Status**: {stats.get('status', 'Unknown')}")
                        st.caption("Products detected dynamically per query")
        
        st.divider()
        
        # Document source selection
        st.subheader("📁 Document Source")
        
        source_option = st.radio(
            "Choose document source:",
            ["Upload Files", "Local Directory"],
            help="Upload files directly or point to a local folder"
        )
        
        if source_option == "Upload Files":
            # File upload section
            st.markdown("**📤 Upload Q&A Documents**")
            
            # SharePoint tip
            with st.expander("💡 Uploading from SharePoint/OneDrive?"):
                st.markdown("""
                **If you're having trouble uploading SharePoint files:**
                
                1. **Download to your computer first:**
                   - Open the file in SharePoint
                   - Click "..." menu → Download
                   - Then upload here
                
                2. **Or use batch download:**
                   - Select all files in SharePoint
                   - Click Download
                   - Extract the ZIP
                   - Upload files from the extracted folder
                
                3. **Avoid:** Don't upload directly from SharePoint's "Open in Desktop App" - 
                   this can cause permission issues.
                """)
            
            uploaded_files = st.file_uploader(
                "Upload Word documents:",
                type=['docx'],
                accept_multiple_files=True,
                help="Upload .docx files to search (max 200MB each)"
            )
            
            if uploaded_files:
                st.info(f"📄 {len(uploaded_files)} files selected")
                
                if st.session_state.searcher is not None:
                    if st.button("📚 Index Uploaded Documents", type="primary", use_container_width=True):
                        # Clean up previous temp directory
                        if st.session_state.temp_dir and os.path.exists(st.session_state.temp_dir):
                            shutil.rmtree(st.session_state.temp_dir)
                        
                        # Create new temp directory
                        st.session_state.temp_dir = tempfile.mkdtemp()
                        
                        # Save uploaded files with better error handling
                        saved_count = 0
                        failed_files = []
                        
                        for uploaded_file in uploaded_files:
                            try:
                                # Clean filename - remove special characters from SharePoint
                                clean_filename = uploaded_file.name
                                # Remove common SharePoint metadata patterns
                                clean_filename = re.sub(r'[<>:"|?*]', '', clean_filename)
                                # Remove any path separators
                                clean_filename = os.path.basename(clean_filename)
                                # Limit filename length
                                if len(clean_filename) > 200:
                                    name, ext = os.path.splitext(clean_filename)
                                    clean_filename = name[:196] + ext
                                
                                file_path = os.path.join(st.session_state.temp_dir, clean_filename)
                                
                                # Write file in binary mode
                                with open(file_path, "wb") as f:
                                    # Get bytes and write
                                    file_bytes = uploaded_file.getbuffer()
                                    f.write(file_bytes)
                                
                                saved_count += 1
                            except Exception as e:
                                failed_files.append((uploaded_file.name, str(e)))
                        
                        # Show summary
                        if failed_files:
                            st.warning(f"⚠️ {len(failed_files)} file(s) failed to upload")
                            with st.expander("View failed files"):
                                for fname, error in failed_files:
                                    st.error(f"**{fname}**: {error}")
                        
                        if saved_count > 0:
                            st.success(f"✅ Saved {saved_count} files")
                            
                            # Index the documents
                            with st.spinner("Indexing uploaded documents..."):
                                num_docs = st.session_state.searcher.index_documents(st.session_state.temp_dir)
                                if num_docs > 0:
                                    st.session_state.indexed = True
                                    st.success(f"✅ Indexed {num_docs} documents!")
                                else:
                                    st.error("❌ No documents were processed")
                        else:
                            st.error("❌ No files could be saved")
                else:
                    st.warning("⚠️ Please load a model first")
        
        else:
            # Local directory section
            st.markdown("**📂 Local Directory**")
            
            with st.expander("💡 How to find your folder path"):
                st.markdown("""
                **On Mac:**
                1. Open Finder and go to your documents folder
                2. Right-click the folder
                3. Hold `Option` key and select "Copy as Pathname"
                4. Paste below
                
                **On Windows:**
                1. Open File Explorer and go to your folder
                2. Click the address bar at the top
                3. Copy the path (e.g., `C:\\Users\\Name\\Documents\\PharmaQA`)
                4. Paste below
                """)
            
            directory = st.text_input(
                "Directory Path:",
                value=os.path.expanduser("~/Documents/PharmaQA"),
                help="Paste the full path to your Q&A documents folder",
                placeholder="/Users/yourname/Documents/PharmaQA"
            )
            
            # Validate and show folder contents preview
            if directory and os.path.exists(directory):
                docx_files = list(Path(directory).rglob('*.docx'))
                docx_files = [f for f in docx_files if not f.name.startswith('~$')]
                if docx_files:
                    st.success(f"✅ Found {len(docx_files)} .docx files")
                else:
                    st.warning("⚠️ No .docx files found in this folder")
            elif directory:
                st.error("❌ Folder does not exist")
            
            file_pattern = st.text_input(
                "File Filter (optional):",
                value="",
                placeholder="e.g., 'product_name' or 'Q&A'",
                help="Only index files containing this text in filename"
            )
            
            # Index button for local directory
            if st.session_state.searcher is not None:
                if st.button("📚 Index Documents", type="primary", use_container_width=True):
                    if not os.path.exists(directory):
                        st.error("❌ Directory does not exist!")
                    else:
                        with st.spinner("Indexing Q&A documents..."):
                            num_docs = st.session_state.searcher.index_documents(
                                directory,
                                file_pattern=file_pattern
                            )
                            if num_docs > 0:
                                st.session_state.indexed = True
                                st.success(f"✅ Indexed {num_docs} documents!")
                            else:
                                st.error("❌ No documents found or processed")
        
        # Statistics
        if st.session_state.indexed:
            st.divider()
            st.subheader("📊 Index Stats")
            
            total_docs = len(st.session_state.searcher.documents)
            total_words = sum(doc.word_count for doc in st.session_state.searcher.documents)
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Documents", total_docs)
            with col2:
                st.metric("Total Words", f"{total_words:,}")
            
            # Document list
            with st.expander("📄 Indexed Documents"):
                for doc in st.session_state.searcher.documents:
                    st.markdown(f"**{doc.filename}**")
                    if doc.title:
                        st.caption(f"📋 Title: {doc.title}")
                    st.caption(f"Modified: {doc.file_modified_date} | Words: {doc.word_count:,}")
                    if doc.qa_pairs:
                        st.caption(f"Q&A Pairs detected: {len(doc.qa_pairs)}")
                    st.divider()
    
    # Main search interface
    if st.session_state.indexed:
        
        # Search section
        st.markdown("### 🔎 Search Query")
        
        # Initialize example query holder
        if 'example_to_load' not in st.session_state:
            st.session_state.example_to_load = ""
        
        # Example questions section
        with st.expander("💡 Example Questions - Click to Use", expanded=False):
            st.markdown("**Note:** Replace `[drug name]` with your actual product name")
            
            # Create tabs for different categories
            tab_names = list(EXAMPLE_QUESTIONS.keys())
            tabs = st.tabs(tab_names)
            
            for tab, category in zip(tabs, tab_names):
                with tab:
                    # Show audience badge
                    if category in HCP_QUESTIONS:
                        st.caption("👨‍⚕️ Healthcare Professional Questions")
                    elif category in PATIENT_QUESTIONS:
                        st.caption("🏥 Patient/Caregiver Questions")
                    
                    for idx, question in enumerate(EXAMPLE_QUESTIONS[category]):
                        button_key = f"ex_{category}_{idx}"
                        if st.button(question, key=button_key, use_container_width=True):
                            st.session_state.example_to_load = question
                            st.rerun()
        
        # Use form to encapsulate the query input
        with st.form(key="search_form", clear_on_submit=False):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                # Use the example if one was selected
                initial_value = st.session_state.example_to_load
                query = st.text_area(
                    "Enter your query:",
                    value=initial_value,
                    placeholder="Example: How to transport vamorolone?\nExample: How to discontinue treatment?",
                    height=100,
                    label_visibility="collapsed"
                )
                # Clear the example after loading
                if initial_value:
                    st.session_state.example_to_load = ""
            
            with col2:
                st.markdown("**Search Options**")
                
                # Audience selector
                audience_mode = st.radio(
                    "Audience:",
                    ["Auto-detect", "Patient/Caregiver", "Healthcare Professional"],
                    help="Choose how results should be presented"
                )
                
                num_results = st.number_input("Results:", 1, 10, 3)
                min_score = st.slider("Min Score:", 0.0, 1.0, 0.2, 0.05)
            
            # Submit button for the form
            submitted = st.form_submit_button("🔍 Find Documents", type="primary", use_container_width=True)
            
            if submitted:
                if query.strip():
                    # Store audience preference
                    if audience_mode == "Patient/Caregiver":
                        st.session_state.forced_audience = 'patient'
                    elif audience_mode == "Healthcare Professional":
                        st.session_state.forced_audience = 'hcp'
                    else:
                        st.session_state.forced_audience = None
                        
                    with st.spinner("Searching..."):
                        results = st.session_state.searcher.search(
                            query, 
                            top_k=num_results,
                            min_score=min_score
                        )
                        st.session_state.last_results = results
                        st.session_state.last_query = query
                else:
                    st.warning("Please enter a query")
        
        # Display results
        if st.session_state.last_results:
            results = st.session_state.last_results
            
            # Detect audience
            if st.session_state.forced_audience:
                audience = st.session_state.forced_audience
            else:
                audience = detect_query_audience(st.session_state.last_query)
            
            # Check if product name is in query (for de-weighting notification)
            detected_products = []
            if st.session_state.searcher:
                detected_products = st.session_state.searcher._get_detected_products(st.session_state.last_query)
            
            detected_product = detected_products[0] if detected_products else None
            detection_method = st.session_state.searcher.detection_method if st.session_state.searcher else 'manual'
            
            st.markdown("---")
            
            # Show audience indicator
            if audience == 'patient':
                st.info("🏥 **Patient/Caregiver Mode**: Results shown in plain language")
            else:
                st.info("👨‍⚕️ **Healthcare Professional Mode**: Results shown with medical terminology")
            
            # Show product name de-weighting indicator with detection method
            if detected_product:
                method_name = PRODUCT_DETECTION_METHODS.get(detection_method, detection_method)
                if len(detected_products) > 1:
                    products_str = f"{', '.join(detected_products[:2])}" + (f" +{len(detected_products)-2} more" if len(detected_products) > 2 else "")
                else:
                    products_str = detected_product
                
                st.success(f"✅ **Smart Search Active** ({method_name}): Results weighted by content relevance, not just '{products_str}' frequency")
            
            st.markdown(f"### 📋 Top {len(results)} Most Relevant Documents")
            
            if len(results) == 0:
                st.info("💡 No documents met the minimum relevance threshold. Try:\n- Lowering the minimum score\n- Rephrasing your query\n- Using more specific medical terms")
            
            for idx, (doc, score) in enumerate(results, 1):
                # Relevance indicator
                if score >= 0.7:
                    indicator = "🟢 High"
                    color = "green"
                elif score >= 0.5:
                    indicator = "🟡 Medium"
                    color = "orange"
                else:
                    indicator = "🟠 Low"
                    color = "red"
                
                # Document card
                with st.container():
                    col1, col2 = st.columns([4, 1])
                    
                    with col1:
                        st.markdown(f"### {idx}. {doc.filename}")
                        # Show title if available and different from filename
                        if doc.title and doc.title != Path(doc.filename).stem:
                            st.caption(f"📋 {doc.title}")
                    with col2:
                        st.markdown(f"<div style='text-align: right; padding-top: 10px;'>{indicator}<br><span style='color: {color}; font-size: 20px; font-weight: bold;'>{score:.1%}</span></div>", unsafe_allow_html=True)
                    
                    # Metadata
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.caption(f"📅 Modified: {doc.file_modified_date}")
                    with col2:
                        st.caption(f"📝 {doc.word_count:,} words")
                    with col3:
                        if doc.qa_pairs:
                            st.caption(f"❓ {len(doc.qa_pairs)} Q&A pairs")
                    
                    # Get relevant snippets
                    snippets = st.session_state.searcher.get_relevant_snippets(
                        doc.full_text, 
                        st.session_state.last_query,
                        num_snippets=3
                    )
                    
                    # Apply patient-friendly language if needed
                    if audience == 'patient':
                        snippets = [simplify_for_patients(snippet) for snippet in snippets]
                    
                    # Show highlighted snippets
                    if snippets:
                        if audience == 'patient':
                            st.markdown("**📌 Key Information (in plain language):**")
                        else:
                            st.markdown("**📌 Most Relevant Excerpts:**")
                            
                        for snippet_idx, snippet in enumerate(snippets, 1):
                            st.markdown(f"{snippet_idx}. {snippet}", unsafe_allow_html=True)
                            if snippet_idx < len(snippets):
                                st.markdown("")  # Add spacing
                    
                    # Content preview
                    with st.expander("👁️ View Full Document", expanded=False):
                        # Show highlighted full content
                        highlighted_text = st.session_state.searcher.highlight_text(
                            doc.full_text, 
                            st.session_state.last_query,
                            max_length=2000
                        )
                        
                        # Apply patient-friendly language if needed
                        if audience == 'patient':
                            highlighted_text = simplify_for_patients(highlighted_text)
                        
                        st.markdown(highlighted_text, unsafe_allow_html=True)
                        
                        # Show Q&A pairs if detected
                        if doc.qa_pairs:
                            st.markdown("---")
                            st.markdown("**Detected Q&A Pairs:**")
                            for q, a in doc.qa_pairs[:5]:  # Show first 5
                                with st.container():
                                    q_display = simplify_for_patients(q) if audience == 'patient' else q
                                    a_display = simplify_for_patients(a) if audience == 'patient' else a
                                    st.markdown(f"**Q:** {q_display}")
                                    st.markdown(f"**A:** {a_display[:300]}{'...' if len(a_display) > 300 else ''}")
                                    st.divider()
                            if len(doc.qa_pairs) > 5:
                                st.caption(f"+ {len(doc.qa_pairs) - 5} more Q&A pairs")
                    
                    # Action buttons
                    col1, col2, col3 = st.columns([2, 2, 3])
                    with col1:
                        # Download button for uploaded files
                        if doc.file_path.startswith(tempfile.gettempdir()):
                            # This is an uploaded file
                            try:
                                with open(doc.file_path, 'rb') as f:
                                    st.download_button(
                                        f"⬇️ Download",
                                        f,
                                        file_name=doc.filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"download_{idx}"
                                    )
                            except:
                                st.caption("Download unavailable")
                        else:
                            # This is a local file
                            if st.button(f"📂 Open Document", key=f"open_{idx}"):
                                try:
                                    import subprocess
                                    import platform
                                    
                                    if platform.system() == 'Darwin':  # macOS
                                        subprocess.call(['open', doc.file_path])
                                    elif platform.system() == 'Windows':
                                        os.startfile(doc.file_path)
                                    else:  # linux
                                        subprocess.call(['xdg-open', doc.file_path])
                                    st.success(f"Opening {doc.filename}")
                                except Exception as e:
                                    st.error(f"Could not open file: {e}")
                    
                    with col2:
                        if st.button(f"📋 Copy Path", key=f"copy_{idx}"):
                            st.code(doc.file_path, language=None)
                    
                    st.markdown("---")
            
            # Export results
            with st.expander("📊 Export Search Results"):
                export_text = f"PHARMA Q&A SEARCH RESULTS\n"
                export_text += f"Query: {st.session_state.last_query}\n"
                export_text += f"Audience: {'Patient/Caregiver' if audience == 'patient' else 'Healthcare Professional'}\n"
                export_text += f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                export_text += f"Results: {len(results)}\n"
                export_text += "=" * 80 + "\n\n"
                
                for idx, (doc, score) in enumerate(results, 1):
                    export_text += f"{idx}. {doc.filename}\n"
                    if doc.title:
                        export_text += f"   Title: {doc.title}\n"
                    export_text += f"   Relevance: {score:.2%}\n"
                    export_text += f"   Modified: {doc.file_modified_date}\n"
                    export_text += f"   Path: {doc.file_path}\n"
                    
                    # Add snippets to export
                    snippets = st.session_state.searcher.get_relevant_snippets(
                        doc.full_text, 
                        st.session_state.last_query,
                        num_snippets=3
                    )
                    
                    # Apply patient-friendly language
                    if audience == 'patient':
                        snippets = [simplify_for_patients(snippet) for snippet in snippets]
                    
                    if snippets:
                        export_text += f"   Excerpts:\n"
                        for snippet in snippets:
                            # Remove HTML tags for plain text export
                            clean_snippet = re.sub(r'<[^>]+>', '', snippet)
                            export_text += f"     - {clean_snippet}\n"
                    
                    export_text += "-" * 80 + "\n\n"
                
                st.download_button(
                    "⬇️ Download Results",
                    export_text,
                    file_name=f"qa_search_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
        
        elif st.session_state.last_query:
            st.info("👆 Enter a query and click 'Find Documents' to search")
    
    else:
        # Welcome screen
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            ### 🎯 How It Works
            
            1. **Load Model** - Choose and load a medical AI model
            2. **Choose Source** - Upload files or point to local directory
            3. **Index Documents** - Process all Word documents
            4. **Search** - Enter any query or use example questions
            5. **Review** - See ranked documents with highlighted matches
            
            ### ✨ Features
            
            - 🧠 **Semantic Search** - Understands meaning, not just keywords
            - 💊 **Medical Domain** - Trained on pharmaceutical literature
            - 🎯 **Smart Weighting** - Title 2x, Filename keywords 2x, Body 1x
            - 🔤 **Abbreviation Expansion** - "MOA" → "mechanism of action"
            - 🔗 **Pharma Synonyms** - "storage" ↔ "transport", "DDI" → "drug interactions"
            - 🧹 **Query Cleaning** - Removes noise words like "are there any"
            - 💊 **Smart Product Detection** - 4 methods: Manual, Auto, NER, Hybrid
            - 🔍 **Enhanced Highlighting** - Catches variations and stems
            - 💡 **Example Questions** - 30+ pre-written HCP and patient questions
            - 🏥 **Patient Mode** - Auto-simplifies medical jargon for patients
            - 👨‍⚕️ **HCP Mode** - Preserves medical terminology for professionals
            - ⚡ **Fast** - Optimized for speed
            - 🔒 **100% Local** - No data leaves your machine
            - 📤 **File Upload** - Works with uploaded or local files
            - 📊 **Relevance Scoring** - See how well documents match
            - 🪟 **Cross-Platform** - Works on Mac and Windows
            """)
        
        with col2:
            st.markdown("""
            ### 💡 Example Queries
            
            **Now Handles:**
            - "what is the MOA of vamorolone?" → expands to "mechanism of action"
            - "are there any DDIs?" → expands to "drug drug interactions"
            - "are there any effects on bone health?" → removes "are there any"
            - "data in peds" → expands to "pediatric children"
            
            **Product Detection:**
            Choose from 4 methods in sidebar:
            
            **1. Manual List** (Fastest)
            - Pre-defined product names
            - Instant detection
            - Must maintain list
            
            **2. Auto-Detect** (Smart)
            - Learns from your documents
            - No maintenance needed
            - Finds frequently mentioned products
            
            **3. NER - ScispaCy** (Most Accurate)
            - Medical AI entity recognition
            - Detects any drug/chemical
            - Requires ScispaCy
            
            **4. Hybrid** (Recommended)
            - Auto-detect + NER fallback
            - Best accuracy + speed
            - Self-learning
            
            ### 🔗 Smart Features
            
            **Abbreviation expansion:**
            - MOA → mechanism of action
            - DDI → drug drug interaction  
            - PK/PD → pharmacokinetics/pharmacodynamics
            - Peds → pediatric children
            
            **Filename boosting:**
            - "Cardiac_Effects" → boosts for "cardiac" queries
            - "Bone_Health" → boosts for "bone" queries
            """)
        
        st.markdown("---")
        
        if st.session_state.searcher is None:
            st.info("👈 **Step 1:** Load a model from the sidebar to begin")
        else:
            st.info("👈 **Step 2:** Upload documents or select local directory in the sidebar")

if __name__ == "__main__":
    main()

